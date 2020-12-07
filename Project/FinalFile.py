import openpyxl as op
import coloredlogs,logging
import sqlite3
import PyPDF2
import time
import docx
import sys
from openpyxl.styles import Font, Fill
from selenium import webdriver
from docx2pdf import convert

coloredlogs.install(level="ERROR")
coloredlogs.install(level="WARNING")

def open_workbook():
    print("-----------------------------------------------------------------------------------------------")
    xlsxFile=input("Enter an Excel workbook Name:")

    try:
        print("OPENING WORKBOOK Employee_Data......")
        wb = op.load_workbook(xlsxFile) # Open workbook Employee Data i.e Employee_Data.xlsx
        try:
            sheetname=input("Enter Worksheet Name: ")
            sheet = wb.get_sheet_by_name(sheetname) # Gets Sheet by sheetname i.e Employee data
            print()
            sheet.title="emp_data" # Sets title to emp_data
            print()
            mysheet=wb.active
            print("Active sheet {}".format(mysheet))
            print()
            print("There are {} rows and {} columns in Employee Data".format(sheet.max_row,sheet.max_column)) # Highest row and column of the sheet
            sheet["F1"]="Job Hours" # Column heading changes from Job_Time to Job Hours
            return sheet,wb
        except KeyError: # Catches Exception when sheet doesnt exist
            logging.error(" Worksheet does not exist.")
            exit()

    except Exception as e:
        logging.error("File not found. Check the filename") # Catches Exception when file doesn&#39;t exist and displays error message
        logging.warning(e)
        exit()


def create_table():
    print("-----------------------------------------------------------------------------------------------")
    print("* Create Table in Database")
    conn = sqlite3.connect(ConStr) #Connect to Connection String
    cr=conn.cursor()

    cr.execute("""CREATE TABLE Employee
    (Emp_id INTEGER PRIMARY KEY, Emp_name TEXT, DOB TEXT,
    Job_Cat INTEGER, Salary INTEGER, Job_Hour TEXT,Project INTEGER,Prev_Exp
    INTEGER,Gender text);""") # Create table Employee

    print("Table is Created")
    conn.commit() #Commit Changes
    conn.close() #Close Connection

def insert_into_table(sheet):
    print("-----------------------------------------------------------------------------------------------")
    print("* Perform Insertion by inserting all Data\&#39;s of Excel to Database ")
    conn = sqlite3.connect(ConStr) #Connect to Connection String
    cr=conn.cursor()
    for i in range(2, sheet.max_row): # Iterate for maximum rows
        for j in range(1,sheet.max_column+1): # Iterate for maximum columns
            Emp_id,Emp_name,DOB,Job_Cat,Salary=sheet.cell(row=i,column=1).value,sheet.cell(row=i,column=2).value,sheet.cell(row=i,column=3).value,sheet.cell(row=i,column=4).value,sheet.cell(row=i,column=5).value
            Job_Time,Project,Prev_Exp,Gender=sheet.cell(row=i,column=6).value,sheet.cell(row=i,column=7).value,sheet.cell(row=i,column=8).value,sheet.cell(row=i,column=9).value
            cr.execute("""INSERT OR IGNORE INTO Employee
            (Emp_id,Emp_name,DOB,Job_Cat,Salary,Job_Hour,Project,Prev_Exp,Gender)
            VALUES ( ?, ?, ?, ?, ?, ? ,?,?,?)""",(Emp_id,Emp_name,DOB,Job_Cat,int(Salary),Job_Time,Project,Prev_Exp,Gender))
            # Inserts all Excel data into Database ExcelData
    print("Data Inserted into table successfully")
    conn.commit() #Commit Changes
    conn.close() #Close Connection

def delete_rows():
    print("------------------------------------------------------------------------------------------------")
    print("* Perform Deletion on Table by deleting records with Emp id greater then 480")
    conn = sqlite3.connect(ConStr)
    cr=conn.cursor()
    cr.execute("""DELETE FROM Employee WHERE Emp_id>480;""") #As there are 490 rows in table, all rows which are greater then 480 are deleted
    print(" Deletion successfull")
    conn.commit()
    conn.close()

def update_table():
    print("------------------------------------------------------------------------------------------------")
    print("* Perform Updation for Table Employee with Experience greater then 20 and with names Starting from S")
    conn = sqlite3.connect(ConStr)
    cr=conn.cursor()
    cr.execute("""UPDATE Employee
    SET Salary = 80000, Prev_Exp= Prev_Exp +1
    WHERE Emp_name LIKE "s%" and Prev_Exp>20;""") #Updates the table by setting salary as 80000, and incrementing there Previous Experience
                                                    #for employee with name that starts and ends with s and a respectively and Experience&gt;20
    print("Table Updated successfully")
    conn.commit()
    conn.close()

def select_cols_based_on_project():
    print("------------------------------------------------------------------------------------------------")
    print("* Query to get Employee\'s Id,Name, Age,Salary,Projects,Experience From table Employee based on their Project,Prev Experience and Salary")
    print()
    conn = sqlite3.connect(ConStr)
    cr=conn.cursor()
    REC=cr.execute("""SELECT Emp_id,Emp_name,Salary,
            case
        when date(dob, '+' ||
            strftime('%Y', 'now') - strftime('%Y', dob) ||
            ' years') >= date('now')
        then strftime('%Y', 'now') - strftime('%Y', dob)
        else strftime('%Y', 'now') - strftime('%Y', dob) - 1
    end
    as Age,
	project,Prev_Exp
    FROM Employee WHERE PROJECT >150 and Prev_Exp >=15 and Salary>=80000  ORDER BY Project Desc;""") #Selects Emp_id,name,salary,age,project,Prev_Exp from Database
    print('Query Executed successfully')
    return REC
    conn.commit()
    conn.close()

def display_records(rec):
    print('---------------------------------------------------------------------------------------------------------------------------------')
    print('Display All Records from Query')
    conn = sqlite3.connect(ConStr)
    cr=conn.cursor()
    print()
    print('EMPLOYEE DETAILS BASED ON QUERY')
    print()
    print('Employee Id, '+'Employee Name, '+'Salary, '+'Age, '+'Project, '+'Prev Experience')  #Display Records stored in  database
    records=rec.fetchall()									#Fetch all query data and store in records variable
    for r in range(len(records)):
        print(str(records[r][0])+'         '+str(records[r][1])+'          '+str(records[r][2])+'     '+str(records[r][3])+'    '+str(records[r][4])+'       '+str(records[r][5]))

    return records
    conn.commit()
    conn.close()

def storing_newdata_excel(records,wb):

    print('---------------------------------------------------------------------------------------------------------------------------------')
    print('Storing Database Query details to Excel')
    try:
        mysht=input(' Enter New Sheet name: ') #Checks for validity
        wb.create_sheet(title=mysht)
        sheet=wb.get_sheet_by_name(mysht)
        sheet['A1'].font= Font(size=14, italic=True,bold=True)      #Styling Excel Sheets
        sheet['A1'] = 'EMPLOYEE\'S DATA FROM DATABASE'

        sheet['A2'],sheet['B2'],sheet['C2'],sheet['D2'],sheet['E2'],sheet['F2']='Emp Id ','Emp Name ','Salary ','Age ','Project ','Prev Experience'           #Set cell values

        i=0
        for rowNum in range(3,len(records)+3):				#Iterate to insert into Excel cells from 3rd row
            sheet.cell(row=rowNum,column=1).value,sheet.cell(row=rowNum,column=2).value,sheet.cell(row=rowNum,column=3).value=records[i][0],records[i][1],records[i][2] #Store record values
            sheet.cell(row=rowNum,column=4).value,sheet.cell(row=rowNum,column=5).value,sheet.cell(row=rowNum,column=6).value=records[i][3],records[i][4],records[i][5] #To excel sheet cells
            i+=1

        try:
            filesave=input("Enter Name of Excel File you want to store data as: ")
            print('  Data Stored Successfully ')
            wb.save(filesave)
            return mysht,filesave #Save WORKBOOK with given filename
        except Exception as e:
            logging.warning("Enter Valid Name of File ")

    except Exception as e:
        logging.error(e)

def convert_xlsx_to_docx(title,filesave):
    print("----------------------------------------------------------------------------------------------")
    print(" Converting Excel to docx ")

    doc = docx.Document() #Create a Document
    try:
        wb = op.load_workbook(filesave) #Load workbook NewEmp_Data which is result of query
        worksheet = wb.get_sheet_by_name(title)
        doc.add_paragraph("Employee Detail of Database query",'Title').add_run()


        ws_range = worksheet.iter_cols(min_col=1, max_col=6, min_row=2,max_row=11, values_only=False) #Produces cells from the worksheet, by column.
        #Specify the iteration range using indices of rows and columns.

        for row in ws_range:
            s = ''
            for cell in row:
                if cell.value is None:
                    s += '' * 11

                else:
                    s += str(cell.value).rjust(10) + '' #Store data in s variable
            doc.add_paragraph(s)  #Add stored data to word document as paragraph
            doc.add_paragraph("____________________________________________________________________________________________________________________")
        doc.add_picture("emp.jpg", width=docx.shared.Inches(5),height=docx.shared.Cm(10)) #Add Picture to word docs
        try:
            wordFile=input("Enter Word Document Name to save Excel data to Word: ") #Should have .docx
            doc.save(wordFile+'.docx') #Save word Document

            print("SAVED SUCCESSFULLY TO WORD")
        except:
            logging.warning("Enter valid file name")

    except Exception as e:
        logging.error("File does not exist or does not contain any data")


def convert_docx_to_pdf():
    print("------------------------------------------------------------------------------------------------")
    print('Converting Word Document to PDF')
    try:
        wordD=input('Enter Word Document to Convert to Pdf: ')
        pdfFile=input("Enter name you who like your pdf to have: ")
        convert(wordD+'.docx','C:/Users/srish/Desktop/PythonProject/{}.pdf'.format(pdfFile))
        #docx2pdf converts docs to pdf
    except Exception as e:
        logging.error('ERROR: {}'.format(e))



def pdfmerge():
	
	print('---------------------------------------------------------------------------------------------------------------------------------')
	print('Merging Files')
	try:
		F1=input('Enter first Pdf file name to merge: ')
		file1=open(F1+'.pdf', 'rb') #Opens pdf in read binary mode
		FirstFile = PyPDF2.PdfFileReader(file1) #Reads all contents of infile1

		F2=input('Enter second Pdf file name to merge: ')
		file2=open(F2+'.pdf', 'rb')
		SecondFile = PyPDF2.PdfFileReader(file2)
		FinalFile = PyPDF2.PdfFileWriter()          #PdfFileWriter file holds the final file

		for i in range(FirstFile.getNumPages()):   #Iterates till highest Number of pages in PDF1
			FinalFile.addPage(FirstFile.getPage(i))  #ADDS every page of firstfile to FinalFile

		for i in range(SecondFile.getNumPages()):   #Iterates till highest Number of pages in PDF2
			FinalFile.addPage(SecondFile.getPage(i))  #ADDS every page of secondfile to FinalFile


		Out=input('Enter Name you want to keep for your Merged file:')
		outfile = open(Out+'.pdf', 'wb') #Open final pdf in write binary mode
		FinalFile.write(outfile) #write each file
		print('Successfully Merged PDFs')

		file2.close()
		file1.close()                               # Closes all files that Was opened
		outfile.close()
	except Exception as e:
        logging.warning('Filename does not Exist')
        logging.error(e)

def mail_automation():

    print('---------------------------------------------------------------------------------------------------------------------------------')
    print('Initializing Mail Automation ')
    try:
        print('Enter Reciever\'s Email ID: ')
        reciever=input()
        print('Enter Message to be sent')
        messageToSend=input()
        messageToSend2=input()
        print()

        browser=webdriver.Chrome()                                                      #Open Web Driver for Chrome
        browser.get("http://gmail.com")                                                 #Get url to open gmail
        email=browser.find_element_by_id('identifierId')

        email.send_keys("SendersId@gmail.com")                                   #Sends senders mail id as key to browser
        browser.find_element_by_id("identifierNext").click()                            #Clicks next button by finding element by id
        time.sleep(3)                                                                   #waits for 3 seconds
        browser.find_element_by_name("password").send_keys('Mypassword')                #Sends users password
        browser.find_element_by_id("passwordNext").click()
        time.sleep(3)                                                                   #waits for 3 seconds

        browser.find_element_by_xpath('//*[@id=":iw"]/div/div').click()                 #Finds element by path
        time.sleep(5)                                                                   #waits for 5 seconds
        to=browser.find_element_by_name("to")                                           #finds element by name
        to.send_keys(reciever)                                    #Send Recievers emailid
        subject=browser.find_element_by_name("subjectbox")
        subject.send_keys(" Python Automation Project")                                 #Sends Subject as key
        MSG=browser.find_element_by_xpath('//*[@id=":ou"]')
        MSG.send_keys(messageToSend)                 #Sends Body Messages
        MSG.send_keys(messageToSend2)
        time.sleep(2)
        attach=browser.find_element_by_xpath('//*[@id=":nh"]').click()
        time.sleep(5)                                                                   #waits for 5 seconds
        browser.close()

    except Exception as e:
        logging.error(e)                                                           #Close browser



if __name__=="__main__":

    print()									#Displays what project does
    print("Mini Project to Demonstrate Working of following Concepts: ")
    print("*********************************************************")

    print("1. Excel: Reading and Writing To Excel files with openpyxl module")
    print("2. Database: Perform CRUD operation by connecting to Sqlite3")
    print("3. Word: Convert Excel file of with query results to Word with pydocx module")
    print("4. PDF: Convert Word to Pdf and Merge different PDF files with PyPDF2 module")
    print("5. Selenium: Automate mail with Selenium webdriver")

    print("***********************************************************")
    print(" Let\'s Get Started ")
    print()
    ConStr='EmpExcelData.db'
    print()

    print("Here are your Choices:: Choose anyone :)")
    print("Choice 1: Opening Excel file ")


    print("Choice 2: Performing CRUD Operations on Database")
    print("Choice 3: Display Database Query Result")
    print("Choice 4: Store Results to Excel")
    print("Choice 5: Convert Excel to Word Document")
    print("Choice 6: Convert Your Word Document to PDF")
    print("Choice 7: PDF Merger")
    print("Choice 8: Automate your Mail")

    while True:
        print()
        print("--------------------------------------------------------------------------------------------")
        print("If want to terminate Enter 0")
        ch=int(input(("Enter your Choice: ")))
        if ch==0:
            print(" Thank You :)")
            sys.exit(0)
        else:
            if ch==1:
                sheet,wb=open_workbook() # Function open_workbook is Called
            elif ch==2:
                create_table()
                insert_into_table(sheet)
                delete_rows()
                update_table()


            elif ch==3:
                REC=select_cols_based_on_project()
                records=display_records(REC)
            elif ch==4:
                t,file=storing_newdata_excel(records,wb)
            elif ch==5:
                convert_xlsx_to_docx(t,file)
            elif ch==6:
                convert_docx_to_pdf()
            elif ch==7:
                pdfmerge()
            elif ch==8:
                mail_automation()
            else:
                print("Choice does not exist !!")


