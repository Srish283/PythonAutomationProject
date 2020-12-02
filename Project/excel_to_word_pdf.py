import openpyxl
import docx
from docx2pdf import convert

doc = docx.Document()                  #Create a Document

workbook = openpyxl.load_workbook('NewEmp_Data.xlsx', data_only=True) #Load workbook NewEmp_Data which is result of query
worksheet = workbook.get_sheet_by_name('Updated Employee Data')




doc.add_paragraph('Employee Detail of Database query', 'Title').add_run()

ws_range = worksheet.iter_cols(min_col=1, max_col=6, min_row=2, max_row=11, values_only=False)

for row in ws_range:
    s = ''
    for cell in row:
        if cell.value is None:
            s += ' ' * 11

        else:
            s += str(cell.value).rjust(10) + ' '

    doc.add_paragraph(s)
    doc.add_paragraph('---------------------------------------------------------------------------------------------------------------------------')
doc.save('Excel_to_Doc.docx')
convert('Excel_to_Doc.docx','C:/Users/srish/Desktop/PythonProject/Docx_to_Pdf.pdf')
